"""
HireMind AI - Presentation Script & Project Explanation Generator
Generates a fully formatted Word (.docx) document.
Run: python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helper: set paragraph shading (background color) ────────────────────────
def shade_paragraph(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def set_cell_bg(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

# ─── Helper: set page margins ─────────────────────────────────────────────────
def set_margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

# ─── Helper: add a heading with custom style ──────────────────────────────────
def add_main_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)   # deep blue
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x27, 0x5E, 0xBF)
        run.font.bold = True
    return p

def add_section_heading(doc, text):
    """Prominent section banner."""
    p = doc.add_paragraph()
    shade_paragraph(p, '1E3A8A')
    run = p.add_run(f"  {text}  ")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_body(doc, text, bullet=False, indent=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bold_body(doc, label, value, bullet=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    run1 = p.add_run(label + ": ")
    run1.font.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    run2 = p.add_run(value)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_info_box(doc, text, color_hex='EFF6FF', border_color='1E3A8A'):
    """Lightly shaded info box."""
    p = doc.add_paragraph()
    shade_paragraph(p, color_hex)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    run.font.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Inches(0.15)
    return p

def add_script_line(doc, speaker, text):
    """Format a script speaker line."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{speaker}]: ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)  # purple
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_divider(doc):
    p = doc.add_paragraph("─" * 90)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def add_tech_table(doc, rows):
    """Adds a 2-column formatted table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Layer / Category"
    hdr[1].text = "Technologies"
    set_cell_bg(hdr[0], '1E3A8A')
    set_cell_bg(hdr[1], '1E3A8A')
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(11)
    for layer, tech in rows:
        row = table.add_row().cells
        row[0].text = layer
        row[1].text = tech
        # Alternate row shading
        set_cell_bg(row[0], 'EFF6FF')
        set_cell_bg(row[1], 'F8FAFC')
    doc.add_paragraph()

def add_module_box(doc, title, purpose, key_functions, tech_used):
    """Box for each module."""
    p = doc.add_paragraph()
    shade_paragraph(p, 'DBEAFE')
    r = p.add_run(f"  MODULE: {title}")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

    add_bold_body(doc, "Purpose", purpose)
    add_bold_body(doc, "Key Functions", key_functions)
    add_bold_body(doc, "Technologies", tech_used)
    doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ═════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()
    set_margins(doc)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ─── COVER PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(title_p, '1E3A8A')
    tr = title_p.add_run("  HireMind AI  ")
    tr.font.size = Pt(32)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(sub_p, '1D4ED8')
    sr = sub_p.add_run("An Autonomous HR Operating System for Intelligent Recruitment")
    sr.font.size = Pt(15)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()
    info_lines = [
        ("Competition", "CALIBRE 2k26 — National Level Project Poster Competition"),
        ("Track",       "Artificial Intelligence & Machine Learning"),
        ("Team",        "Kartik Gupta, Bhavesh Chaudhary, Durga Muhto, Sudiksha Singh, Komal Yadav"),
        ("Institution", "Thakur Shyamanarayan Engineering College, Mumbai, India"),
        ("Date",        "18th April 2026"),
    ]
    for label, value in info_lines:
        add_bold_body(doc, label, value)

    doc.add_page_break()

    # ─── TABLE OF CONTENTS ────────────────────────────────────────────────────
    add_section_heading(doc, "TABLE OF CONTENTS")
    toc_items = [
        "PART 1: DETAILED PROJECT EXPLANATION",
        "   1.1  Project Overview & Problem Statement",
        "   1.2  System Architecture",
        "   1.3  Technology Stack Explained",
        "   1.4  Module-by-Module Deep Dive",
        "         A. Data Ingestion (data_ingestion.py)",
        "         B. Text Processing & NLP Parsing (text_processing.py)",
        "         C. Embeddings & FAISS Vector Store (embeddings.py)",
        "         D. Semantic Similarity Matching (similarity.py)",
        "         E. ATS Scoring Engine (ats_scorer.py)",
        "         F. Fraud Detection (fraud_detector.py)",
        "         G. Identity Verification (identity_verification.py)",
        "         H. Webcam Proctoring Monitor (webcam_monitor.py)",
        "         I. Interview Engine — Dual Mode (interview_engine.py)",
        "         J. Voice Interview System (voice_interview.py)",
        "         K. LLM Analysis (llm_analysis.py)",
        "         L. Recommendation Engine (recommendation_engine.py)",
        "   1.5  End-to-End Data Flow",
        "   1.6  Database Layer",
        "",
        "PART 2: POSTER PRESENTATION SCRIPT",
        "   2.1  Introduction",
        "   2.2  Objectives",
        "   2.3  Methodology",
        "   2.4  Results & Discussion",
        "   2.5  Observations",
        "   2.6  Conclusion",
        "   2.7  Q&A Preparation Tips",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        if item:
            run = p.add_run(item)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    #  PART 1 — DETAILED PROJECT EXPLANATION
    # ═══════════════════════════════════════════════════════════════════════════

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(title_p, '0F172A')
    t = title_p.add_run("PART 1: DETAILED PROJECT EXPLANATION")
    t.font.size = Pt(18)
    t.font.bold = True
    t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

    # ── 1.1 Project Overview ──────────────────────────────────────────────────
    add_section_heading(doc, "1.1  PROJECT OVERVIEW & PROBLEM STATEMENT")

    add_body(doc,
        "The modern recruitment process is fundamentally broken. Human recruiters spend 23 hours "
        "screening resumes for a single hire, unconscious bias affects up to 74% of hiring decisions, "
        "and the rise of generative AI has created an 'Integrity Crisis' — candidates now use AI tools "
        "to generate deceptive, keyword-stuffed resumes that fool traditional ATS systems."
    )
    add_body(doc,
        "HireMind AI is a complete, multi-layered autonomous HR operating system built entirely in Python "
        "using Streamlit as the interface layer. It transforms every stage of recruitment — from the "
        "moment a resume is uploaded to the moment a hire decision is made — into an intelligent, "
        "data-driven, and bias-resistant workflow."
    )

    add_subsection_heading(doc, "Core Problems Solved:")
    problems = [
        "Slow manual screening → Automated in seconds using NLP and vector search",
        "Keyword-stuffed/AI-generated fraud resumes → Detected using 4 independent heuristic checks",
        "Impersonation risk → Face recognition compares resume photo against live webcam capture",
        "Biased interviews → AI-generated, candidate-specific questions ensure structured evaluation",
        "No cheating during interviews → Continuous MediaPipe webcam proctoring with violation tracking",
        "Admin overload → Single dashboard ranks all candidates with combined scores and one-click hiring",
    ]
    for p in problems:
        add_body(doc, p, bullet=True)

    doc.add_paragraph()

    # ── 1.2 System Architecture ───────────────────────────────────────────────
    add_section_heading(doc, "1.2  SYSTEM ARCHITECTURE")

    add_body(doc,
        "HireMind AI follows a layered microservice-style architecture with five distinct horizontal "
        "layers. Each layer has independent modules that communicate through structured Python dictionaries "
        "(not HTTP calls), making the system fast, testable, and easy to debug."
    )

    add_subsection_heading(doc, "Architecture Layers:")

    layers = [
        ("Layer 1 — User Interface",
         "Streamlit (Python) web dashboard with custom CSS glassmorphism styling. Two user "
         "roles: Recruiter (admin panel, job posting, candidate management) and Candidate "
         "(resume upload, verification, interview participation)."),
        ("Layer 2 — API Gateway & Authentication",
         "Session-state management via Streamlit's built-in session. OTP-based email "
         "authentication (Resend API). SQLite-backed auth module (auth_db.py) for secure "
         "user account management, password hashing, and session tokens."),
        ("Layer 3 — Backend Processing Services",
         "Four core services: (1) Candidate Service — processes uploads and triggers the full "
         "analysis pipeline. (2) Job Management Service — handles job description storage and "
         "linkage. (3) Workflow Orchestrator — app.py coordinates all modules in sequence. "
         "(4) ATS Integration — connects ATS results to the FAISS semantic index."),
        ("Layer 4 — AI Modules & Processing",
         "Eight specialized AI modules: Resume Analysis, Semantic Matching, ATS Scoring, Fraud "
         "Detection, Interview Engine, Voice Interview, Webcam Proctoring, and LLM Analysis. "
         "Each module is self-contained with its own logic, thresholds, and output schema."),
        ("Layer 5 — Security & Verification",
         "Identity Verification (face_recognition + dlib), OTP Gateway, Webcam Proctoring "
         "(MediaPipe), and fraud detection heuristics form the security perimeter."),
        ("Layer 6 — Data Storage",
         "Three-tier storage: SQLite (relational — user accounts, scores, interview records), "
         "Local filesystem (raw and processed resume files), FAISS vector store (semantic "
         "embeddings for all candidate resumes — persistence via faiss.index + metadata.pkl)."),
    ]

    for title, desc in layers:
        add_bold_body(doc, title, "")
        add_body(doc, desc, indent=True)

    doc.add_paragraph()

    # ── 1.3 Technology Stack ──────────────────────────────────────────────────
    add_section_heading(doc, "1.3  TECHNOLOGY STACK — EXPLAINED")

    add_body(doc, "Every technology choice in HireMind AI was made deliberately. Here is the full stack with rationale:")
    doc.add_paragraph()

    tech_rows = [
        ("Frontend / UI",         "Streamlit 1.56 — Python-native web framework for data-science apps. "
                                  "Custom CSS3 glassmorphism theming applied via st.markdown()."),
        ("LLM / AI Backbone",     "Google Gemini 2.0 Flash (primary) — Used for: personalized interview "
                                  "question generation, voice answer evaluation, deep resume-JD analysis, "
                                  "and fraud validation. Fallback: OpenAI GPT-4o-mini."),
        ("NLP — Text Parsing",    "spaCy en_core_web_md — Named Entity Recognition (PERSON, GPE/LOC) "
                                  "for candidate name and location extraction. Medium model chosen "
                                  "for balance of speed and accuracy."),
        ("NLP — Skill Matching",  "Custom taxonomy of 200+ skills across 12 domains. Case-insensitive "
                                  "regex + substring matching against extracted resume sections."),
        ("Semantic Matching",     "Sentence Transformers (all-MiniLM-L6-v2, 384-dim) — Converts resume "
                                  "chunks and job descriptions into dense semantic vectors. Chosen for "
                                  "its speed and accuracy on short text paragraphs."),
        ("Vector Database",       "FAISS IndexFlatIP (Meta AI) — Inner Product index with L2-normalized "
                                  "vectors gives cosine similarity directly. Supports incremental "
                                  "candidate addition without full rebuild."),
        ("Text Chunking",         "LangChain RecursiveCharacterTextSplitter — Chunk size 400 chars, "
                                  "80 char overlap. Preserves semantic boundaries at sentence/paragraph "
                                  "level before embedding."),
        ("Fuzzy Matching",        "RapidFuzz — Token sort ratio scoring for interview answer evaluation. "
                                  "Handles word-order variations better than exact string matching."),
        ("Speech-to-Text",        "AssemblyAI Universal-3-pro (REST API) — 3-step flow: audio upload, "
                                  "transcript request, polling. 120-second timeout. High accuracy "
                                  "across accents."),
        ("Text-to-Speech",        "gTTS (Google Text-to-Speech) — Converts interview questions to MP3 "
                                  "audio bytes for playback. Creates conversational voice interview UX."),
        ("Face Detection",        "MediaPipe Face Detection — Runs on every webcam frame during "
                                  "interview. CLAHE brightness normalization handles low-light. "
                                  "Detects 0, 1, or multiple faces in real time."),
        ("Face Recognition",      "face_recognition (dlib HOG + CNN encodings) — Compares resume "
                                  "profile photo (extracted from PDF) against live webcam snapshot with "
                                  "0.6 Euclidean distance threshold."),
        ("Webcam Streaming",      "streamlit-webrtc + av (PyAV) — WebRTC-based video stream with "
                                  "STUN server (stun.l.google.com). Runs in SENDRECV mode."),
        ("PDF Extraction",        "PyPDF2 (primary) → pdfminer.six (fallback) → OCR placeholder. "
                                  "Image extraction uses PyPDF2 XObject parsing for profile photos."),
        ("Database",              "SQLite3 (via Python stdlib) — Three tables: users, candidates, "
                                  "interview_answers. No external server needed."),
        ("Email Service",         "Resend API — Sends OTP verification emails and automated "
                                  "hiring/rejection decision emails to candidates."),
        ("Environment",           "python-dotenv — Loads GEMINI_API_KEY, ASSEMBLYAI_API_KEY, "
                                  "RESEND_API_KEY from .env file securely."),
    ]

    add_tech_table(doc, tech_rows)

    # ── 1.4 Module Deep Dive ──────────────────────────────────────────────────
    add_section_heading(doc, "1.4  MODULE-BY-MODULE DEEP DIVE")

    # MODULE A
    add_subsection_heading(doc, "MODULE A — Data Ingestion (data_ingestion.py)")
    add_body(doc,
        "This is the entry point for every candidate's resume. When a recruiter or candidate uploads "
        "a file, this module handles the complete ingestion pipeline."
    )
    add_bold_body(doc, "Supported Formats", "PDF, DOCX, TXT")
    add_bold_body(doc, "Extraction Strategy",
        "PDF: PyPDF2 (primary) → pdfminer.six (fallback) → OCR placeholder (last resort). "
        "DOCX: python-docx paragraph extraction. TXT: UTF-8 decode with error replacement.")
    add_subsection_heading(doc, "Key Functions:")
    funcs_a = [
        ("save_uploaded_file()",    "Saves Streamlit UploadedFile to data/resumes/raw/ with a "
                                    "timestamped filename to prevent collisions. Sanitizes filename "
                                    "using regex."),
        ("_extract_images_from_pdf()", "Extracts the largest image on Page 1 of the PDF "
                                    "(heuristic: width × height). Converts to RGB JPEG and "
                                    "returns base64-encoded bytes for identity verification."),
        ("clean_text()",            "7-step text normalizer: UTF-8 re-encode → remove control chars "
                                    "→ normalize unicode dashes/quotes/bullets → collapse spaces "
                                    "→ collapse blank lines → strip lines → final strip."),
        ("process_resume()",        "Full pipeline orchestrator: save → extract → clean → save. "
                                    "Returns a structured dict with cleaned_text, word_count, "
                                    "char_count, and profile_photo (base64)."),
    ]
    for name, desc in funcs_a:
        add_bold_body(doc, name, desc, bullet=True)

    doc.add_paragraph()

    # MODULE B
    add_subsection_heading(doc, "MODULE B — Text Processing & NLP Parsing (text_processing.py)")
    add_body(doc,
        "Transforms raw cleaned text into a structured JSON dictionary — the core data object "
        "that flows through every downstream module."
    )
    add_subsection_heading(doc, "Pipeline Steps:")
    steps_b = [
        "Step 1 — Advanced cleaning: removes boilerplate ('References on request', 'I hereby declare'), "
        "URLs, and punctuation-only lines using pre-compiled regex patterns.",
        "Step 2 — Section segmentation: builds a composite regex from 8 section heading taxonomies "
        "(education, experience, skills, projects, summary, certifications, achievements, languages). "
        "Identifies heading boundaries and assigns text blocks to each section.",
        "Step 3 — Entity extraction: spaCy NER on first 3000 chars for PERSON (name) and GPE/LOC "
        "(location). Regex global search on full text for email and phone (with +91 normalization). "
        "Handles PDF fragmentation artifacts ('B h a v e s h' → 'Bhavesh') via regex pre-processing.",
        "Step 4 — Skill extraction: case-insensitive matching against a 200+ skill taxonomy across "
        "12 domains. Only extracts from the 'skills' section (strict) to avoid false positives.",
        "Step 5 — JSON assembly: combines all extracted fields into a single structured dict with "
        "keys: name, email, phone, location, skills, summary, education, experience, projects, "
        "certifications, achievements, languages, profile_photo.",
    ]
    for s in steps_b:
        add_body(doc, s, bullet=True)

    doc.add_paragraph()

    # MODULE C
    add_subsection_heading(doc, "MODULE C — Embeddings & FAISS Vector Store (embeddings.py)")
    add_body(doc,
        "Converts the structured resume JSON into high-dimensional vector representations "
        "stored in a persistent FAISS index for semantic retrieval."
    )
    steps_c = [
        ("prepare_text_chunks()", "Maps each resume section (skills, experience, education, projects, "
                                  "summary, certifications, achievements) into a labeled text block "
                                  "(e.g., 'Work Experience:\\n...') for context-aware embedding. "
                                  "Also creates an overview chunk combining name + top skills + summary."),
        ("chunk_text()",          "Long segments are split with LangChain RecursiveCharacterTextSplitter "
                                  "(chunk_size=400, overlap=80). Short segments (<400 chars) are kept intact. "
                                  "Priority separators: double newline → single newline → sentence → comma → space."),
        ("generate_embeddings()", "Encodes all chunks using all-MiniLM-L6-v2 with normalize_embeddings=True "
                                  "(L2-normalization). Returns float32 numpy array of shape (n_chunks, 384)."),
        ("create_faiss_index()",  "Creates FAISS IndexFlatIP — Inner Product index. Since vectors are "
                                  "L2-normalized, IP equals cosine similarity. The index stores all "
                                  "candidate chunk vectors together."),
        ("save_vector_store()",   "Persists faiss.index (binary) and metadata.pkl (chunk text + section + "
                                  "name). Also writes metadata.json for human inspection."),
        ("Incremental Updates",   "When a new candidate is added, existing chunks for that candidate are "
                                  "removed (name-based filter), new chunks appended, and the full index "
                                  "is rebuilt. Handles deduplication correctly."),
    ]
    for name, desc in steps_c:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE D
    add_subsection_heading(doc, "MODULE D — Semantic Similarity Matching (similarity.py)")
    add_body(doc,
        "Performs job-description-to-resume matching using vector similarity search. "
        "This is the core ranking engine behind the recruiting dashboard."
    )
    steps_d = [
        ("process_job_description()", "Cleans JD text and extracts required skills using the same "
                                      "taxonomy as the resume parser."),
        ("generate_jd_embedding()",   "Embeds the cleaned JD using the same model+normalization "
                                      "as resumes to ensure compatible vector space."),
        ("search_similar_resumes()",  "FAISS IndexFlatIP.search() returns Top-K nearest chunks "
                                      "with cosine similarity scores. k defaults to 20."),
        ("aggregate_by_candidate()",  "Groups chunk-level results by candidate name. Computes: "
                                      "max_score, mean_score, combined_score = 0.6 × max + 0.4 × mean. "
                                      "Deduplicates near-identical chunks (hash of first 150 chars). "
                                      "Applies 0.40 threshold to filter weak matches."),
        ("filter_by_hard_skills()",   "Case-insensitive set intersection between candidate skills "
                                      "and JD skills. Returns matched/missing skills + match_ratio."),
        ("rank_candidates()",         "Final score = (semantic_score × 0.65) + (skill_match_ratio × 0.35). "
                                      "Sorted descending. Rewards depth of semantic match while "
                                      "penalizing missing critical skills."),
    ]
    for name, desc in steps_d:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE E
    add_subsection_heading(doc, "MODULE E — ATS Scoring Engine (ats_scorer.py)")
    add_body(doc,
        "A fully deterministic, explainable ATS (Applicant Tracking System) evaluator. "
        "No AI. No external APIs. Purely rule-based and transparent."
    )
    add_bold_body(doc, "Scoring Formula",
        "ATS Score = (0.60 × Section Score) + (0.40 × Structure Score)")
    steps_e = [
        ("check_section_presence()", "Checks for 4 required sections: Skills, Experience, Education, Projects. "
                                     "Each present section = 25 points. Section Score out of 100."),
        ("validate_experience_structure()", "3 checks: (1) Date presence — minimum 2 date patterns detected "
                                            "using 3 regex variants (Month Year, YYYY-YYYY, standalone year). "
                                            "(2) Chronological order — years are generally decreasing (most "
                                            "recent first). (3) Content depth — experience section ≥30 words. "
                                            "Structure Score: has_dates (+50) + is_chronological (+30) + depth (+20)."),
        ("generate_ats_feedback()", "Produces prioritized, actionable suggestions: missing sections first, "
                                    "then date issues, then ordering issues, then score-tier advice "
                                    "(Poor/Fair/Good/Excellent)."),
        ("Grade Assignment",        "≥80 = Excellent, ≥60 = Good, ≥40 = Fair, <40 = Poor."),
    ]
    for name, desc in steps_e:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE F
    add_subsection_heading(doc, "MODULE F — Fraud Detection (fraud_detector.py)")
    add_body(doc,
        "Detects resume manipulation, exaggeration, and ATS gaming using four independent "
        "deterministic checks. No embeddings. No AI required. Fully explainable."
    )
    steps_f = [
        ("1. Skill Inconsistency Check",   "Cross-references skills listed in the 'Skills' section "
                                           "against the body text (experience + projects + education). "
                                           "Any skill with zero word-boundary matches in the body is "
                                           "flagged as undemonstrated. Skips tokens shorter than 3 chars "
                                           "(single-letter languages like 'C' or 'R')."),
        ("2. Keyword Stuffing Detection",   "Tokenizes raw text, filters stopwords, counts word frequencies. "
                                           "Flags any non-stopword appearing >8 times AND with density >3% "
                                           "of total resume text. Reports the word, count, and density %."),
        ("3. Hidden Text Detection",        "Three heuristics: (a) Raw word count ≥2.5× parsed word count "
                                           "suggests invisible white text. (b) Regex finds 5+ consecutive "
                                           "repetitions of the same word. (c) Lines with 20+ words and "
                                           "zero punctuation indicate keyword dump paragraphs."),
        ("4. Temporal Consistency Check",   "Extracts date ranges from experience text using 3 regex patterns. "
                                           "Detects: impossible timelines (end < start), future start dates, "
                                           "suspiciously long tenures (>15 years), and overlapping employment "
                                           "periods (>3 months overlap between two roles)."),
        ("Risk Scoring",                    "Each flag carries a severity weight: Hidden Text / Timeline Error = 3, "
                                           "Keyword Stuffing / Overlapping = 2, Skill Inconsistency = 1. "
                                           "Raw score = weight_total × 15, capped at 100. "
                                           "Risk levels: Low (≤20), Medium (≤55), High (>55)."),
        ("Optional LLM Validation",         "On Medium/High risk, Gemini is invoked with detected flags + "
                                           "resume excerpt to provide a 2-3 sentence professional verdict "
                                           "distinguishing true fraud from false positives."),
    ]
    for name, desc in steps_f:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE G
    add_subsection_heading(doc, "MODULE G — Identity Verification (identity_verification.py)")
    add_body(doc,
        "Three-layer identity cross-validation to ensure the person submitting the resume "
        "is who they claim to be."
    )
    steps_g = [
        ("Name Verification (fuzzy)",  "Python difflib.SequenceMatcher compares parsed name (from resume NER) "
                                       "against the name entered by the user in the registration form. "
                                       "Threshold: 85% similarity ratio. Returns match status and score."),
        ("Email Verification (strict)", "Exact case-insensitive string match between resume-extracted email "
                                        "and registered email."),
        ("Phone Verification (suffix)", "Compares last 10 digits of both phone numbers to handle country "
                                        "code variations (+91 vs 91 vs blank)."),
        ("OTP Gateway",                 "generate_otp() creates a cryptographically random 6-digit code. "
                                        "verify_otp_logic() checks: OTP match + expiry timestamp (via time.time()). "
                                        "OTP sent via Resend API email."),
        ("Face Recognition",            "verify_identity(): Decodes base64 resume photo → face_recognition.face_encodings() "
                                        "→ Decodes live webcam snapshot → face_encodings() → compare_faces() with "
                                        "tolerance=0.6 → face_distance() converted to similarity score (0–1)."),
    ]
    for name, desc in steps_g:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE H
    add_subsection_heading(doc, "MODULE H — Webcam Proctoring Monitor (webcam_monitor.py)")
    add_body(doc,
        "Provides real-time continuous monitoring during the interview session to prevent "
        "impersonation and unauthorized assistance."
    )
    steps_h = [
        ("streamlit-webrtc",          "WebRTC SENDRECV mode with Google STUN server. Captures video at "
                                      "640×480 resolution at 20fps. FaceDetectionProcessor runs in a "
                                      "background thread."),
        ("MediaPipe Face Detection",  "model_selection=0 (short-range, optimized for faces <2m away). "
                                      "min_detection_confidence=0.5. Processes every 0.5 seconds for efficiency."),
        ("CLAHE Normalization",       "Contrast Limited Adaptive Histogram Equalization on the L-channel "
                                      "of LAB color space. Handles poor lighting conditions during remote interviews."),
        ("Violation System",          "2-strike system: (1) No face detected for >2 seconds → Warning 1/2. "
                                      "(2) Multiple faces detected → instant Warning. At 2 violations → "
                                      "interview terminated and session reset. Status buffer prevents false "
                                      "positives from frame flicker."),
        ("Visual Overlay",            "Bounding boxes drawn on detected faces (green for 1, red for multiple). "
                                      "Status label ('FACE OK' / 'NO FACE' / 'X FACES DETECTED') overlaid "
                                      "on the video stream with high-contrast background."),
    ]
    for name, desc in steps_h:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE I
    add_subsection_heading(doc, "MODULE I — Interview Engine — Dual Mode (interview_engine.py)")
    add_body(doc,
        "The core text-based interview system supports two distinct evaluation modes: "
        "Normal (deterministic, rule-based) and AI Mode (Gemini-powered)."
    )
    add_bold_body(doc, "Stage 1 — Normal Mode",
        "7 questions: 3 Coding, 2 Technical, 2 Behavioral. Questions drawn from a structured question bank "
        "(question_bank.py) with predefined keywords and expected answers.")
    steps_i_normal = [
        ("Coding Evaluation",    "Composite score: (code-similarity × 0.4) + (test_case_pass_rate × 0.4) + "
                                 "(keyword_boost × 0.2). Safe exec() runs actual test cases against the "
                                 "candidate's submitted code. Security filter blocks os, sys, subprocess, "
                                 "open(), __import__. RapidFuzz ratio against expected answer."),
        ("Text Evaluation",      "RapidFuzz token_sort_ratio (handles word-order reordering) = 70% weight. "
                                 "Keyword presence boost = 30% weight. 'Keyword boost' = (matched_keywords / "
                                 "total_keywords) × 20."),
        ("Score Tiers",          "80-100% composite → 9-10 marks, 60-79% → 7-8 marks, 40-59% → 5-6 marks, "
                                 "below 40% → 0-4 marks (max 10 per question)."),
    ]
    for name, desc in steps_i_normal:
        add_bold_body(doc, name, desc, bullet=True)

    add_bold_body(doc, "Stage 1 — AI Mode",
        "generate_ai_questions(): Sends candidate skills to Gemini to generate 7 personalized questions "
        "(3 coding + 2 technical + 2 behavioral). evaluate_with_ai(): Gemini scores each answer 0-10 "
        "with a 2-sentence evaluation. generate_deep_ai_report(): Gemini analyzes all 7 responses "
        "together and returns a hireability verdict (Strong Hire / Hire / No Hire).")
    doc.add_paragraph()

    # MODULE J
    add_subsection_heading(doc, "MODULE J — Voice Interview System (voice_interview.py & voice_ui.py)")
    add_body(doc,
        "Stage 2 of the interview — a voice-based conversational interview that evaluates communication "
        "clarity along with technical depth. Questions are read aloud; answers are spoken and automatically transcribed."
    )
    steps_j = [
        ("Question Generation",   "generate_voice_questions(): Sends resume data (name, top 8 skills, "
                                  "top 3 projects) + JD summary (first 600 chars) to Gemini. "
                                  "Generates exactly 3 questions: project_based, skill_based, scenario_based."),
        ("Text-to-Speech",        "text_to_speech(): gTTS converts question text to MP3 bytes in English. "
                                  "Played back via st.audio() in the browser."),
        ("Speech-to-Text",        "transcribe_audio(): 3-step AssemblyAI REST flow: "
                                  "(1) POST /v2/upload → upload raw audio bytes. "
                                  "(2) POST /v2/transcript with universal-3-pro model + language_code=en. "
                                  "(3) GET /v2/transcript/{id} polling every 3 seconds, 120-second deadline."),
        ("Answer Evaluation",     "evaluate_voice_answer_async(): Non-blocking Gemini evaluation in a "
                                  "background thread. Evaluates: correctness, depth, relevance, communication. "
                                  "Score 0-10 with 2-3 sentence reasoning."),
        ("Combined Report",       "generate_combined_report(): Weights Stage 1 (60%) + Stage 2 (40%). "
                                  "Verdict: Strong Hire (≥75%), Hire (≥55%), No Hire (<55%). "
                                  "Final Gemini analysis synthesizes both stages into key strengths, "
                                  "weaknesses, skill gaps, and hiring recommendation."),
    ]
    for name, desc in steps_j:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE K
    add_subsection_heading(doc, "MODULE K — LLM Analysis (llm_analysis.py)")
    add_body(doc,
        "Provides deep qualitative AI analysis of a resume against a job description — going "
        "far beyond keyword matching to semantic and cultural fit reasoning."
    )
    steps_k = [
        ("build_analysis_prompt()", "Constructs a 6-section expert prompt: Match Analysis → Strengths → "
                                    "Weaknesses/Gaps → Soft Skills Analysis → Cultural Fit → Recommendations. "
                                    "Injects parsed resume context (skills, education, experience, projects) "
                                    "for higher quality output."),
        ("call_llm()",              "Supports Gemini (gemini-2.0-flash default) and OpenAI (gpt-4o-mini). "
                                    "Gemini: retry logic for 429 (quota) and 404 (model not found) errors. "
                                    "Temperature=0.3, max_tokens=4096."),
        ("parse_llm_output()",      "Strips markdown code block wrappers. JSON.loads() direct parse. "
                                    "Fallback: regex JSON extraction from response. Type validation and "
                                    "coercion for all 10 expected fields. match_percentage clamped to 0-100."),
        ("Output Schema",           "10-field JSON: match_percentage, match_summary, strengths (list), "
                                    "weaknesses (list), matched_skills, missing_skills, "
                                    "soft_skills_analysis, cultural_fit, recommendations, interview_questions."),
    ]
    for name, desc in steps_k:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # MODULE L
    add_subsection_heading(doc, "MODULE L — Recommendation Engine (recommendation_engine.py)")
    add_body(doc,
        "A knowledge-graph-driven (no AI required) career development system that generates "
        "personalized learning paths for candidates with skill gaps."
    )
    steps_l = [
        ("identify_skill_gaps()",          "Set-based intersection with substring matching (e.g., 'machine learning' "
                                           "matches 'ml'). Returns matched_skills, missing_skills, extra_skills, match_rate."),
        ("SKILL_KNOWLEDGE_GRAPH",          "Hardcoded graph of 40+ skills mapping each to: prerequisites chain, "
                                           "2 recommended courses, 1 project idea, 1 certification. Covers: AI/ML, "
                                           "Web Dev, Cloud/DevOps, Data Engineering, Programming Languages."),
        ("map_skills_to_learning_paths()", "Looks up each missing skill in the knowledge graph. Falls back to "
                                           "generic 'Search Udemy/Coursera' advice for unknown skills."),
        ("suggest_alternative_roles()",    "ROLE_SKILL_MAP matches candidate skills against 12 predefined roles "
                                           "(Data Analyst, ML Engineer, Full Stack Developer, etc.). Returns "
                                           "top 4 roles with ≥30% skill match."),
        ("generate_improvement_summary()", "Produces tiered summary: Strong Match (≥80%), Moderate Match (≥50%), "
                                           "Significant Gap (<50%). Lists top 5 priority skills and recommended domains."),
    ]
    for name, desc in steps_l:
        add_bold_body(doc, name, desc, bullet=True)
    doc.add_paragraph()

    # ── 1.5 End-to-End Data Flow ──────────────────────────────────────────────
    add_section_heading(doc, "1.5  END-TO-END DATA FLOW")

    add_body(doc,
        "The following describes the complete journey of a candidate through HireMind AI from "
        "first upload to final hiring decision:"
    )

    flow_steps = [
        ("STAGE 0 — Registration",
         "Candidate registers on the web dashboard. auth_db.py creates an account with hashed "
         "credentials. OTP sent via Resend API to verify email ownership."),
        ("STAGE 1A — Resume Upload",
         "Candidate uploads PDF/DOCX/TXT. data_ingestion.py: saves raw file → extracts text "
         "(PyPDF2/pdfminer) → extracts profile photo from Page 1 → cleans text → saves processed. "
         "Returns: cleaned_text + profile_photo_base64."),
        ("STAGE 1B — NLP Processing",
         "text_processing.py: cleans text further → segments into sections → extracts entities "
         "(spaCy NER for name/location, regex for email/phone) → extracts skills (taxonomy match). "
         "Returns: structured_json."),
        ("STAGE 1C — Vector Embedding",
         "embeddings.py: prepares labeled text segments → chunks with LangChain splitter → "
         "encodes with Sentence Transformer → adds to FAISS IndexFlatIP → saves faiss.index + metadata.pkl."),
        ("STAGE 1D — Fraud Check",
         "fraud_detector.py: skill inconsistency check → keyword stuffing detection → "
         "hidden text detection → temporal consistency check → compute risk score. "
         "Returns: fraud_risk (Low/Medium/High), risk_score, flags."),
        ("STAGE 1E — ATS Scoring",
         "ats_scorer.py: section presence check → experience structure validation → "
         "weighted score calculation → feedback generation. Returns: ats_score (0-100), grade."),
        ("STAGE 1F — Semantic Matching",
         "similarity.py: JD text → embed → FAISS search → aggregate by candidate → "
         "skill filtering → rank. Returns: ranked list with final_score, semantic_score, skill_score."),
        ("STAGE 1G — LLM Analysis (optional)",
         "llm_analysis.py: builds expert prompt → calls Gemini → parses JSON → validates schema. "
         "Returns: match_percentage, strengths, weaknesses, soft skills analysis, cultural fit, "
         "recommendations, interview questions."),
        ("STAGE 1H — Identity Verification",
         "identity_verification.py: fuzzy name match + exact email/phone match + OTP verification + "
         "face recognition (resume photo vs webcam capture). All 4 checks must pass."),
        ("STAGE 2A — Text Interview",
         "interview_engine.py: select 7 questions (Normal or AI mode) → candidate submits answers → "
         "evaluate_answer() per question → generate_normal_report(). Webcam monitor active throughout."),
        ("STAGE 2B — Voice Interview",
         "voice_interview.py: generate 3 Gemini voice questions → gTTS playback → "
         "candidate records audio → AssemblyAI transcription → Gemini evaluation → "
         "generate_combined_report() (Stage1 60% + Stage2 40%)."),
        ("STAGE 3 — Admin Dashboard",
         "Recruiter views all candidates sorted by combined ATS + interview score. "
         "One-click 'Select' sends Resend API hiring email + marks candidate as hired in SQLite."),
    ]

    for step, desc in flow_steps:
        add_bold_body(doc, step, "")
        add_body(doc, desc, indent=True)

    doc.add_paragraph()

    # ── 1.6 Database ──────────────────────────────────────────────────────────
    add_section_heading(doc, "1.6  DATABASE LAYER")
    add_body(doc,
        "HireMind AI uses SQLite3 (standard library, no server required) for all relational "
        "data storage. The database has three primary tables:"
    )
    db_tables = [
        ("users table (auth_db.py)",
         "Stores: user_id, name, email (hashed), password_hash, role (recruiter/candidate), "
         "otp_code, otp_expiry, created_at."),
        ("candidates table (candidate_db.py)",
         "Stores: candidate_id, user_id (FK), name, email, phone, skills_json, ats_score, "
         "fraud_risk, interview_score (Stage 1), voice_score (Stage 2), combined_score, "
         "hire_status, job_id (FK), resume_path, profile_photo_b64, created_at."),
        ("interview_answers table",
         "Stores: answer_id, candidate_id (FK), question_text, answer_text, score, evaluation, "
         "question_type (coding/technical/behavioral), interview_mode (normal/ai/voice)."),
        ("FAISS Vector Store (non-relational)",
         "faiss.index: binary FAISS index (all resume chunk vectors). "
         "metadata.pkl: Python list of dicts (chunk_text, section, source, name, chunk_id). "
         "metadata.json: human-readable inspection copy."),
    ]
    for name, desc in db_tables:
        add_bold_body(doc, name, desc, bullet=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    #  PART 2 — POSTER PRESENTATION SCRIPT
    # ═══════════════════════════════════════════════════════════════════════════

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(title_p2, '0F172A')
    t2 = title_p2.add_run("PART 2: POSTER PRESENTATION SCRIPT")
    t2.font.size = Pt(18)
    t2.font.bold = True
    t2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

    add_info_box(doc,
        "PRESENTATION GUIDELINES:\n"
        "• Total time: 5–7 minutes for full script. Each section ~50–70 seconds.\n"
        "• Speak confidently and look at the judges, not the poster.\n"
        "• Keywords are in BOLD — emphasize them when speaking.\n"
        "• Pause briefly at section transitions.\n"
        "• [SPEAKER] = You (the presenter)."
    )

    # ── SECTION 0 — Opening Hook ──────────────────────────────────────────────
    add_section_heading(doc, "OPENING HOOK (Before Pointing to Any Section)")

    add_script_line(doc, "SPEAKER",
        "Good morning / afternoon, respected judges and faculty. "
        "My name is [Your Name], and along with my teammates Kartik, Bhavesh, Durga, Sudiksha, and Komal, "
        "I am proud to present HireMind AI — an autonomous HR operating system that completely reimagines "
        "how companies hire talent in the age of artificial intelligence."
    )
    add_script_line(doc, "SPEAKER",
        "We built this system from scratch over several months using Python, Streamlit, and "
        "a combination of NLP, Computer Vision, Speech Processing, and Large Language Models. "
        "Let me walk you through it — section by section."
    )
    add_divider(doc)

    # ── SECTION 1 — Introduction ──────────────────────────────────────────────
    add_section_heading(doc, "SECTION 1 — INTRODUCTION")
    add_info_box(doc, "[Point to the INTRODUCTION box on the poster]")

    add_script_line(doc, "SPEAKER",
        "Let me start with the problem we are solving. In today's digital era, recruitment systems are "
        "often slow, biased, and inefficient. A typical recruiter spends 23 hours screening resumes "
        "for a single position, and despite this effort, unconscious bias still affects hiring decisions."
    )
    add_script_line(doc, "SPEAKER",
        "But recently, a new and more dangerous problem has emerged. The rise of Generative AI — tools "
        "like ChatGPT — has led to what we call an Integrity Crisis. Candidates now generate polished, "
        "keyword-stuffed resumes that look perfect on paper but don't reflect their actual skills. "
        "They pad their skills sections, hide invisible text for ATS manipulation, or even overlap "
        "their employment dates to appear more experienced."
    )
    add_script_line(doc, "SPEAKER",
        "HR teams are struggling to manage large volumes of such applications effectively. "
        "There was a real need for a system that is not just fast, but also intelligent enough to "
        "detect fraud, verify identity, and evaluate candidates fairly — all automatically. "
        "That is exactly what HireMind AI does."
    )
    add_divider(doc)

    # ── SECTION 2 — Objectives ────────────────────────────────────────────────
    add_section_heading(doc, "SECTION 2 — OBJECTIVES")
    add_info_box(doc, "[Point to the OBJECTIVES box on the poster]")

    add_script_line(doc, "SPEAKER",
        "We had three core objectives when building this system."
    )
    add_script_line(doc, "SPEAKER",
        "First — to develop a fully end-to-end automated hiring system with AI integration. "
        "Every single step — from resume upload to the final hire decision — is handled by the system "
        "without requiring manual HR intervention."
    )
    add_script_line(doc, "SPEAKER",
        "Second — to implement semantic matching and fraud detection for accurate candidate evaluation. "
        "We didn't want to just match keywords. We wanted the system to understand context, to know "
        "that 'built a machine learning pipeline' is relevant to a data science role even if the exact "
        "phrase 'data science' isn't present."
    )
    add_script_line(doc, "SPEAKER",
        "And third — to ensure secure and intelligent screening through identity verification and "
        "AI-driven interviews. We verify that the person sitting behind the screen is actually the "
        "person whose resume was submitted, using facial recognition. And during the interview, "
        "our webcam monitoring system ensures no impersonation or cheating occurs."
    )
    add_divider(doc)

    # ── SECTION 3 — Methodology ───────────────────────────────────────────────
    add_section_heading(doc, "SECTION 3 — METHODOLOGY")
    add_info_box(doc, "[Point to the METHODOLOGY box + Figure 1 System Architecture on the poster]")

    add_script_line(doc, "SPEAKER",
        "Now let me explain HOW we built this. Our system works in two main stages."
    )
    add_script_line(doc, "SPEAKER",
        "Stage 1 is what we call Precision Screening. When a candidate uploads their resume, "
        "our data ingestion module extracts the text using PyPDF2 with pdfminer as a fallback. "
        "It also extracts the profile photo embedded in the PDF for identity verification later."
    )
    add_script_line(doc, "SPEAKER",
        "This text is then parsed using spaCy — a natural language processing library — which "
        "identifies the candidate's name, location, and email using Named Entity Recognition. "
        "We then extract their technical skills against a taxonomy of over 200 skills across 12 domains."
    )
    add_script_line(doc, "SPEAKER",
        "This structured data is then converted into 384-dimensional vector embeddings using Sentence "
        "Transformers and stored in a FAISS vector database. When a recruiter enters a job description, "
        "that description is also embedded into the same vector space, and FAISS finds the most "
        "semantically similar resumes — not just keyword-matched ones — within milliseconds."
    )
    add_script_line(doc, "SPEAKER",
        "Simultaneously, our ATS Scoring Engine checks: does the resume have all four required sections — "
        "Skills, Experience, Education, and Projects? Are there proper date ranges? Is the experience "
        "listed in reverse chronological order? The final ATS score is computed as 60% section score "
        "plus 40% structure score."
    )
    add_script_line(doc, "SPEAKER",
        "Before any candidate proceeds further, our Fraud Detection module runs four independent checks: "
        "Skill Inconsistency — are listed skills actually demonstrated anywhere in the resume? "
        "Keyword Stuffing — is any non-stopword appearing more than 8 times? "
        "Hidden Text Detection — is the raw word count disproportionately larger than parsed content? "
        "And Temporal Consistency — are there overlapping employment dates or impossible timelines?"
    )
    add_script_line(doc, "SPEAKER",
        "If a candidate clears screening, they move to Identity Verification. "
        "Our system uses fuzzy name matching, strict email matching, OTP-based email verification, "
        "and facial recognition using the dlib library. We extract the profile photo from the submitted "
        "PDF and compare it with a live webcam capture."
    )
    add_script_line(doc, "SPEAKER",
        "Stage 2 is the AI Interview System. It has two modes."
    )
    add_script_line(doc, "SPEAKER",
        "In text interview mode, Google Gemini generates seven personalized questions based on the "
        "candidate's actual resume — three coding questions, two technical questions, and two behavioral "
        "questions. Coding answers are evaluated by actually executing the code using Python's exec() "
        "function and running real test cases. Text answers are scored using RapidFuzz similarity "
        "scoring against ideal expected answers."
    )
    add_script_line(doc, "SPEAKER",
        "In voice interview mode, Gemini generates three spoken questions — project-based, skill-based, "
        "and scenario-based. The question is read aloud using gTTS. The candidate speaks their answer, "
        "which is recorded and transcribed using AssemblyAI's Universal-3-pro model. Gemini then "
        "evaluates the transcribed answer on accuracy, depth, and communication clarity."
    )
    add_script_line(doc, "SPEAKER",
        "Throughout the entire interview, our webcam monitoring module uses MediaPipe to detect faces "
        "in real time. If the candidate leaves the frame for more than 2 seconds, or if multiple faces "
        "are detected, a violation is registered. Two violations automatically terminate the interview."
    )
    add_divider(doc)

    # ── SECTION 4 — Results & Discussion ─────────────────────────────────────
    add_section_heading(doc, "SECTION 4 — RESULTS & DISCUSSION")
    add_info_box(doc, "[Point to Figure 2 Workflow, Figure 3 Interview with Face Detection, and the System Architecture diagram]")

    add_script_line(doc, "SPEAKER",
        "Now let me walk you through what the system actually looks like in practice."
    )
    add_script_line(doc, "SPEAKER",
        "Looking at Figure 2 — the workflow diagram — you can see the five-step candidate journey: "
        "Resume Upload → Resume Analysis → Candidate Matching → AI Interview → Final Result. "
        "On the right side, you can see the parallel security layers: Fraud Detection, ID Verification, "
        "and AI Interview all running in parallel before the final result is produced."
    )
    add_script_line(doc, "SPEAKER",
        "Looking at Figure 3 — the AI interview system screenshot — you can see the real-time face "
        "detection in action. The green bounding box confirms the system has detected exactly one face, "
        "and the status bar shows 'FACE OK — Monitoring active.' The interview question is displayed "
        "alongside the webcam feed. This is a completely real implementation — not a mockup."
    )
    add_script_line(doc, "SPEAKER",
        "Looking at Figure 1 — the system architecture — you can see all six layers: User Interface "
        "at the top, then API Gateway, then Backend Processing Services, then AI Modules, "
        "then Security and Verification, and finally Data Storage at the bottom with our three-tier "
        "storage: SQLite for relational data, object storage for resume files, and the FAISS vector "
        "database for semantic embeddings."
    )
    add_script_line(doc, "SPEAKER",
        "Our recruiter dashboard shows all candidates ranked by a combined score that weights "
        "ATS compatibility and interview performance. A one-click selection button sends an "
        "automated hiring email via the Resend API and updates the candidate's status in the database."
    )
    add_divider(doc)

    # ── SECTION 5 — Observations ──────────────────────────────────────────────
    add_section_heading(doc, "SECTION 5 — OBSERVATIONS")
    add_info_box(doc, "[Point to the OBSERVATIONS section on the right side of the poster]")

    add_script_line(doc, "SPEAKER",
        "Based on our testing, we made several important observations."
    )
    add_script_line(doc, "SPEAKER",
        "First, semantic matching using FAISS and Sentence Transformers consistently provides more "
        "relevant candidate-job alignment compared to traditional keyword-based systems. A candidate "
        "who built a 'natural language question answering system' correctly ranks higher for an "
        "NLP role than a candidate who just listed 'NLP' as a skill but never demonstrated it."
    )
    add_script_line(doc, "SPEAKER",
        "Second, our fraud detection mechanism successfully identifies manipulated or inconsistent "
        "resume data. In our tests, resumes with keyword stuffing, identical to popular AI-generated "
        "templates, were correctly flagged with Medium to High fraud risk scores."
    )
    add_script_line(doc, "SPEAKER",
        "Third, identity verification using face recognition significantly improves trust and "
        "authenticity in the hiring process. By comparing the photo in the submitted PDF against "
        "a live webcam capture, we can detect cases where someone else applies on another person's behalf."
    )
    add_script_line(doc, "SPEAKER",
        "Fourth, AI-based interviews — both text and voice — ensure consistent and unbiased evaluation "
        "of all candidates. Every candidate is asked questions tailored to their specific resume, "
        "and every answer is evaluated using the same objective criteria."
    )
    add_script_line(doc, "SPEAKER",
        "And fifth, the integration of all these AI modules together results in a dramatically more "
        "reliable and intelligent hiring pipeline. The system reduces manual effort and screening "
        "time significantly, while improving decision quality."
    )
    add_divider(doc)

    # ── SECTION 6 — Conclusion ────────────────────────────────────────────────
    add_section_heading(doc, "SECTION 6 — CONCLUSION")
    add_info_box(doc, "[Point to the CONCLUSIONS box on the poster]")

    add_script_line(doc, "SPEAKER",
        "Let me close by summarizing what HireMind AI achieves."
    )
    add_script_line(doc, "SPEAKER",
        "First and most importantly, HireMind AI successfully automates the entire recruitment process — "
        "from the moment a resume is received to the moment a hiring decision is made. "
        "No manual screening. No bias. No guesswork."
    )
    add_script_line(doc, "SPEAKER",
        "Semantic analysis using vector embeddings and FAISS significantly improves candidate-job "
        "matching accuracy over traditional keyword systems. Fraud detection mechanisms — using 4 "
        "independent rule-based checks — enhance system reliability and protect recruiters from "
        "manipulated resumes. Identity verification using face recognition ensures that only "
        "authentic candidates proceed through the pipeline."
    )
    add_script_line(doc, "SPEAKER",
        "And AI-powered interviews — both text and voice based, with real-time webcam proctoring — "
        "provide a scalable, fair, and unbiased evaluation of every candidate."
    )
    add_script_line(doc, "SPEAKER",
        "HireMind AI is not just a concept. It is a fully implemented, working system. "
        "Every module you see documented here is real code that runs. "
        "We believe this system represents a genuine advancement in intelligent recruitment technology, "
        "and we are excited to demonstrate it live if you would like to see it in action."
    )
    add_script_line(doc, "SPEAKER",
        "Thank you for your time. We are happy to answer any technical questions you may have."
    )
    add_divider(doc)

    # ── SECTION 7 — Q&A Preparation ──────────────────────────────────────────
    add_section_heading(doc, "SECTION 7 — Q&A PREPARATION GUIDE")
    add_body(doc,
        "Be prepared for the following questions from judges. Confident, specific answers will "
        "impress — always cite actual code when possible."
    )

    qas = [
        ("Q: How does FAISS work, and why did you choose it?",
         "FAISS (Facebook AI Similarity Search) is a library that enables efficient similarity search "
         "over high-dimensional vectors. We chose it because it supports Inner Product search (IndexFlatIP), "
         "which with L2-normalized vectors directly gives cosine similarity scores. It's highly optimized "
         "in C++ with a Python wrapper, handles thousands of vectors in milliseconds on a CPU, and "
         "supports persistence to disk. We use it to find the most similar resumes to a job description "
         "in our vector space."),
        ("Q: How accurate is your fraud detection? What are the false positive rates?",
         "Our fraud detection uses 4 independent rule-based checks with weighted severity scoring. "
         "A single isolated skill inconsistency (weight=1) alone will produce a Low risk score (<20), "
         "not triggering action. We require multiple high-severity flags (hidden text weight=3, "
         "timeline error weight=3) to reach High risk. We also allow a 3-month grace period for "
         "overlapping employment to account for job transitions. False positives are minimized by "
         "design — we err on the side of 'review needed' rather than 'automatic rejection.'"),
        ("Q: How does the voice interview work technically?",
         "It uses a 3-component pipeline. First, gTTS converts the question text to an MP3 audio file "
         "which plays in the browser. The candidate records their answer as audio. That audio is "
         "uploaded to AssemblyAI's REST API via a 3-step flow: upload → request transcript → poll. "
         "The AssemblyAI Universal-3-pro model transcribes it. The transcript is then sent to Gemini "
         "with context about the question type, and Gemini evaluates it on 4 criteria."),
        ("Q: What happens if the face recognition library fails?",
         "We handle this gracefully. In identity_verification.py, we use a try/except import block. "
         "If face_recognition is unavailable, HAS_FACE_REC = False and the function returns a "
         "skipped_lib_missing status instead of crashing. Identity verification can still proceed "
         "with the name, email, phone, and OTP checks."),
        ("Q: How does the webcam monitoring prevent cheating?",
         "The webcam monitor uses streamlit-webrtc to stream live video and processes each frame "
         "with MediaPipe Face Detection at 2 FPS (every 0.5 seconds). A 2-second stability buffer "
         "prevents false violations from momentary frame drops. Violations trigger on: no face "
         "detected for >2 continuous seconds, OR multiple faces detected. A 2-strike system gives "
         "one warning before terminating and resetting the interview session."),
        ("Q: What is the difference between Normal Mode and AI Mode interviews?",
         "Normal Mode uses pre-written questions from a curated question bank (question_bank.py) "
         "covering coding, technical, and behavioral questions. Evaluation is fully deterministic "
         "using RapidFuzz similarity and test case execution — no API call required. AI Mode uses "
         "Gemini to generate 7 personalized questions based on the candidate's actual resume skills. "
         "Evaluation is also Gemini-powered, providing nuanced scoring with natural language feedback."),
        ("Q: How do you combine Stage 1 and Stage 2 scores?",
         "The combined report weighs Stage 1 (text interview) at 60% and Stage 2 (voice interview) "
         "at 40%. Formula: combined_pct = (s1_percentage × 0.60) + (s2_percentage × 0.40). "
         "Verdicts: Strong Hire ≥75%, Hire ≥55%, No Hire <55%. A final Gemini deep analysis "
         "also synthesizes both stages into key strengths, weaknesses, skill gaps, and a "
         "hiring recommendation paragraph."),
    ]

    for q, a in qas:
        add_subsection_heading(doc, q)
        add_body(doc, a)
        doc.add_paragraph()

    # ── FINAL NOTE ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(final_p, '1E3A8A')
    fr = final_p.add_run("  HireMind AI — Setting the Benchmark for Intelligent Recruitment  ")
    fr.font.size = Pt(14)
    fr.font.bold = True
    fr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()
    last = doc.add_paragraph()
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = last.add_run(
        "Developed by: Kartik Gupta, Bhavesh Chaudhary, Durga Muhto, Sudiksha Singh, Komal Yadav\n"
        "Thakur Shyamanarayan Engineering College, Mumbai | CALIBRE 2k26"
    )
    lr.font.size = Pt(11)
    lr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    lr.font.italic = True

    # SAVE
    output_path = r"c:\Users\LENOVO\OneDrive\Desktop\HireMind_AI_Presentation_Script.docx"
    doc.save(output_path)
    print(f"\n[SUCCESS] Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Open your document at: {path}")
