"""
HireMind AI - System Architecture Explanation Document Generator
Generates: HireMind_AI_System_Architecture_Explanation.docx
Run: python generate_architecture_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── HELPER UTILITIES ────────────────────────────────────────────────────────

def shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

def layer_banner(doc, layer_num, title, color_hex, text_color=(255,255,255)):
    p = doc.add_paragraph()
    shade_para(p, color_hex)
    r = p.add_run(f"  LAYER {layer_num}  |  {title}  ")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*text_color)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def step_banner(doc, step_num, title, color_hex='1D4ED8'):
    p = doc.add_paragraph()
    shade_para(p, 'DBEAFE')
    r1 = p.add_run(f"  STEP {step_num}  ")
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    r2 = p.add_run(f"  {title}")
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def sub_heading(doc, text, color=(29, 78, 216)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, indent=False, bullet=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p

def bold_kv(doc, key, value, key_color=(30, 58, 138), bullet=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    r1 = p.add_run(key + ": ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(*key_color)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.space_after = Pt(4)
    return p

def iop_table(doc, inp, proc, out):
    """Input / Processing / Output 3-row table."""
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ('INPUT',      inp,  '1E3A8A', 'EFF6FF'),
        ('PROCESSING', proc, '065F46', 'ECFDF5'),
        ('OUTPUT',     out,  '92400E', 'FFFBEB'),
    ]
    for i, (label, content, label_color, content_color) in enumerate(rows_data):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = content
        set_cell_bg(cells[0], label_color)
        set_cell_bg(cells[1], content_color)
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for para in cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def speaking_box(doc, text):
    p = doc.add_paragraph()
    shade_para(p, 'F3F4F6')
    r = p.add_run("SPEAKING SCRIPT:  " + text)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x40, 0x7A)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(14)
    return p

def divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("_" * 95)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ═════════════════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT BUILDER
# ═════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── COVER ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    cp = doc.add_paragraph()
    shade_para(cp, '0F172A')
    cr = cp.add_run("  System Architecture Explanation  ")
    cr.font.size = Pt(28)
    cr.font.bold = True
    cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sp = doc.add_paragraph()
    shade_para(sp, '1E3A8A')
    sr = sp.add_run("  HireMind AI: An Autonomous HR Operating System for Intelligent Recruitment  ")
    sr.font.size = Pt(13)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    bold_kv(doc, "Competition", "CALIBRE 2k26 — National Level Project Poster Competition")
    bold_kv(doc, "Reference",   "Figure 1: System Architecture of HireMind AI (Poster)")
    bold_kv(doc, "Document",    "System Architecture Step-by-Step Explanation")
    bold_kv(doc, "Date",        "18th April, 2026")
    doc.add_page_break()

    # ── OVERVIEW ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    shade_para(p, '1E3A8A')
    r = p.add_run("  OVERVIEW: HOW TO READ THIS DOCUMENT  ")
    r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(8)

    body(doc,
        "Figure 1 in the poster titled 'System Architecture of HireMind AI' shows a 6-layer "
        "horizontal stack diagram. Reading from TOP to BOTTOM, each layer represents a distinct "
        "tier of the system. Inside each layer, there are labelled boxes representing the "
        "specific services or components at that tier. Arrows connect layers to show how data "
        "flows downward through the system."
    )
    body(doc,
        "This document explains every layer and every component in that diagram — in the exact "
        "order they appear — with full technical detail, plus a short speaking version for "
        "presenting during judging."
    )

    p2 = doc.add_paragraph()
    shade_para(p2, 'DBEAFE')
    r2 = p2.add_run(
        "  LAYER SUMMARY (Top to Bottom): "
        "Layer 1: USER INTERFACE  ->  "
        "Layer 2: API GATEWAY & AUTHENTICATION  ->  "
        "Layer 3: BACKEND PROCESSING SERVICES  ->  "
        "Layer 4a: AI MODULES & PROCESSING  |  Layer 4b: SECURITY & VERIFICATION  ->  "
        "Layer 5: DATA STORAGE  "
    )
    r2.font.size = Pt(10); r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after  = Pt(14)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  PART A — LAYER-BY-LAYER EXPLANATION
    # ══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    shade_para(p, '0F172A')
    r = p.add_run("  PART A: LAYER-BY-LAYER ARCHITECTURE EXPLANATION  ")
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── LAYER 1 ───────────────────────────────────────────────────────────────
    layer_banner(doc, 1, "USER INTERFACE", '1E3A8A')

    body(doc,
        "This is the topmost layer — the part of the system that all human users directly interact with. "
        "Looking at the diagram, Layer 1 contains two main entry points: the Web Dashboard (accessed "
        "by Recruiters on the left) and the Resume Upload Portal (accessed by Candidates on the right)."
    )

    sub_heading(doc, "Component 1A — Web Dashboard (Recruiter Side)")
    body(doc,
        "The Web Dashboard is the recruiter's control panel. It is built entirely using Streamlit — "
        "a Python web framework that renders the UI from pure Python scripts. The dashboard provides "
        "the recruiter with: job posting and management capabilities, a ranked candidate table showing "
        "all applicants sorted by their combined score, access to detailed per-candidate reports "
        "(ATS score, fraud risk, interview score, AI hireability verdict), and a one-click 'Select' "
        "action that triggers the automated hiring email and database update."
    )
    bold_kv(doc, "Technology", "Streamlit 1.56 with Custom CSS3 glassmorphism styling")
    bold_kv(doc, "User Role",  "Recruiter (HR Manager / Admin)")
    bold_kv(doc, "Key Actions", "Create job posts, view candidate rankings, send hiring decisions")

    sub_heading(doc, "Component 1B — Resume Upload Portal (Candidate Side)")
    body(doc,
        "The Resume Upload Portal is the candidate-facing interface. Candidates register, verify their "
        "email via OTP, and upload their resume (PDF, DOCX, or TXT). After upload, the portal guides "
        "the candidate through Identity Verification and then into the Interview System. All UI "
        "interactions are handled through Streamlit's session state, which persists user data across "
        "page interactions without a traditional web session cookie."
    )
    bold_kv(doc, "Technology", "Streamlit UploadedFile widget, session_state for flow control")
    bold_kv(doc, "User Role",  "Candidate (Job Applicant)")
    bold_kv(doc, "Key Actions", "Register, OTP verify, upload resume, complete interview")

    divider(doc)

    # ── LAYER 2 ───────────────────────────────────────────────────────────────
    layer_banner(doc, 2, "API GATEWAY & AUTHENTICATION", '1D4ED8')

    body(doc,
        "Layer 2 sits immediately below the UI layer. In the diagram it spans the full width — "
        "this represents the fact that ALL requests from both Recruiters and Candidates must pass "
        "through this layer before any backend processing begins. It handles identity management, "
        "session validation, and secure API routing."
    )

    sub_heading(doc, "Authentication Sub-System (auth_db.py)")
    body(doc,
        "The authentication module manages: new user registration with bcrypt-style password hashing "
        "stored in SQLite, login validation against hashed credentials, OTP generation and expiry "
        "checking using Python's time.time() timestamps, and role-based access control separating "
        "Recruiter vs. Candidate permissions. All data is stored in the 'users' table of the "
        "SQLite database."
    )
    bold_kv(doc, "Technology", "SQLite3, python-dotenv for key management, Resend API for OTP email")
    bold_kv(doc, "Security",   "OTP has time-limited expiry; role checking prevents candidates accessing recruiter panels")
    bold_kv(doc, "Session",    "Streamlit session_state maintains the authenticated user's context across all pages")

    divider(doc)

    # ── LAYER 3 ───────────────────────────────────────────────────────────────
    layer_banner(doc, 3, "BACKEND PROCESSING SERVICES", '0369A1')

    body(doc,
        "Layer 3 is the orchestration layer — it coordinates the workflows initiated by the UI and "
        "delegates work to the AI modules below. Looking at the diagram, Layer 3 contains FOUR "
        "distinct service boxes arranged horizontally: Job Service, Candidate Service, "
        "Job Management Service, and Workflow Orchestrator. On the right side, there is also "
        "an ATS Integration component linked to external Job Boards."
    )

    sub_heading(doc, "Service 3A — Job Service")
    body(doc,
        "Handles everything related to job posting creation and management. When a recruiter creates "
        "a new job, this service stores the job title, description, required skills, and department "
        "in the SQLite 'jobs' table. It also generates the Job ID that is used to link candidates "
        "to a specific opening in the database."
    )
    bold_kv(doc, "Input",  "Recruiter fills in job title, department, description, required skills")
    bold_kv(doc, "Output", "Job record stored in SQLite with a unique Job ID")

    sub_heading(doc, "Service 3B — Candidate Service")
    body(doc,
        "This is the most active service in Layer 3. When a candidate uploads a resume, the "
        "Candidate Service is responsible for triggering the full analysis pipeline in sequence: "
        "it calls data_ingestion.py to extract text, then text_processing.py to build the "
        "structured JSON, then embeddings.py to generate vectors, then runs ATS scoring, "
        "fraud detection, and finally stores all results in the 'candidates' SQLite table. "
        "It acts as the pipeline coordinator."
    )
    bold_kv(doc, "Input",  "Uploaded resume file (PDF/DOCX/TXT) + candidate session info")
    bold_kv(doc, "Output", "Populated candidate record: skills, ATS score, fraud risk, FAISS embedding")

    sub_heading(doc, "Service 3C — Job Management Service")
    body(doc,
        "Handles the matching phase — when a recruiter triggers a search against a job description, "
        "this service coordinates the call to similarity.py (the FAISS semantic search module) "
        "and retrieves the ranked candidate list. It also manages the linkage between each "
        "candidate record and the specific job they applied for."
    )
    bold_kv(doc, "Input",  "Job Description text + Job ID")
    bold_kv(doc, "Output", "Ranked list of candidates sorted by semantic match score")

    sub_heading(doc, "Service 3D — Workflow Orchestrator (app.py)")
    body(doc,
        "This is the master coordinator — the main app.py Streamlit application itself. It renders "
        "each page, reads from session_state to know where the user is in the flow, calls the "
        "appropriate service or module, and passes results between layers. Think of it as the "
        "traffic controller of the entire system. All module calls flow through this orchestrator."
    )
    bold_kv(doc, "Technology", "Streamlit multipage app with st.session_state as the state machine")
    bold_kv(doc, "Role",       "Master coordinator — connects UI to backend to AI modules to storage")

    sub_heading(doc, "Service 3E — ATS Integration & Job Boards (Right side of diagram)")
    body(doc,
        "On the right side of Layer 3, the diagram shows ATS Integration connecting upward to "
        "Job Boards. This represents the system's potential external integration capability — "
        "where the ATS scoring results can be aligned with standard industry ATS formats, "
        "and job postings can be linked to external platforms."
    )

    divider(doc)

    # ── LAYER 4 ───────────────────────────────────────────────────────────────
    layer_banner(doc, 4, "AI MODULES & PROCESSING  |  SECURITY & VERIFICATION", '065F46')

    body(doc,
        "Layer 4 is the intelligence core of HireMind AI. In the diagram, this layer is split "
        "into TWO parallel sections side by side: on the LEFT is 'AI Modules & Processing' "
        "containing three major AI components, and on the RIGHT is 'Security & Verification' "
        "containing three security components. Both sections operate simultaneously during the "
        "screening and interview phases."
    )

    # LEFT SIDE
    p = doc.add_paragraph()
    shade_para(p, 'ECFDF5')
    r = p.add_run("  LEFT SIDE — AI MODULES & PROCESSING  ")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)

    sub_heading(doc, "AI Component 4A — Resume Analysis & Parsing (NLP / NER)")
    body(doc,
        "This component encompasses two modules working together: data_ingestion.py and "
        "text_processing.py. The data ingestion module extracts raw text from the uploaded "
        "resume file. The text processing module then transforms that raw text into a structured "
        "JSON object using three techniques: (1) spaCy Named Entity Recognition identifies the "
        "candidate's name (PERSON entity) and location (GPE/LOC entity) from the first 3000 "
        "characters of the resume. (2) Regex-based pattern matching extracts email addresses "
        "and phone numbers globally across the full text, with automatic +91 Indian number "
        "normalization. (3) Section segmentation splits the resume into labelled sections "
        "(Skills, Experience, Education, Projects, etc.) using a composite heading regex "
        "built from 8 section heading taxonomies. Skills are then extracted from the 'Skills' "
        "section only, matched against a 200+ term taxonomy across 12 domains."
    )
    bold_kv(doc, "Module",      "data_ingestion.py + text_processing.py")
    bold_kv(doc, "Technology",  "spaCy en_core_web_md, PyPDF2, pdfminer.six, regex, python-docx")
    bold_kv(doc, "Input",       "Raw resume file (PDF/DOCX/TXT)")
    bold_kv(doc, "Output",      "Structured JSON: {name, email, phone, location, skills, education, experience, projects}")

    sub_heading(doc, "AI Component 4B — Candidate Evaluation & Matching (Scoring, Skills)")
    body(doc,
        "This component covers three modules: embeddings.py (vector generation), similarity.py "
        "(semantic matching), and ats_scorer.py (ATS scoring). "
        "EMBEDDINGS: The structured JSON sections are converted into labeled text blocks, "
        "chunked using LangChain RecursiveCharacterTextSplitter (400 char chunks, 80 char overlap), "
        "and encoded into 384-dimensional vectors using the all-MiniLM-L6-v2 Sentence Transformer. "
        "These vectors are stored in a FAISS IndexFlatIP for cosine similarity search. "
        "SIMILARITY MATCHING: When a recruiter submits a job description, it is embedded using "
        "the same model. FAISS search returns the 20 most similar resume chunks. These are "
        "aggregated by candidate name, deduplicated, and ranked using the formula: "
        "final_score = (semantic_score × 0.65) + (skill_match_ratio × 0.35). "
        "ATS SCORING: Simultaneously, a rule-based ATS evaluation runs — checking section "
        "presence (60% weight) and experience structure such as date formatting and chronological "
        "ordering (40% weight). The result is a percentage score with a grade (Poor/Fair/Good/Excellent)."
    )
    bold_kv(doc, "Module",      "embeddings.py + similarity.py + ats_scorer.py")
    bold_kv(doc, "Technology",  "Sentence Transformers (all-MiniLM-L6-v2), FAISS IndexFlatIP, LangChain, RapidFuzz")
    bold_kv(doc, "Input",       "Structured resume JSON + Job Description text")
    bold_kv(doc, "Output",      "Ranked candidate list with semantic score, skill score, ATS score, grade")

    sub_heading(doc, "AI Component 4C — AI Interview System (Verbal, Behavioral)")
    body(doc,
        "This component covers interview_engine.py and voice_interview.py. "
        "TEXT INTERVIEW (Stage 1): Google Gemini generates 7 personalized interview questions "
        "based on the candidate's extracted skills — 3 coding, 2 technical, 2 behavioral. "
        "Coding answers are evaluated by executing the submitted code using Python's safe exec() "
        "(with os/sys/subprocess blocked) and running predefined test cases. Text answers are "
        "scored using RapidFuzz token_sort_ratio (70%) + keyword presence boost (30%). "
        "Each answer receives a score out of 10 with a written evaluation. "
        "VOICE INTERVIEW (Stage 2): A separate voice-based interview with 3 Gemini-generated "
        "spoken questions (project-based, skill-based, scenario-based). Questions are converted "
        "to MP3 audio via gTTS and played in the browser. The candidate speaks their answer, "
        "which is transcribed by AssemblyAI (Universal-3-pro) and then evaluated by Gemini on "
        "correctness, depth, relevance and communication. The final combined report weights "
        "Stage 1 at 60% and Stage 2 at 40%."
    )
    bold_kv(doc, "Module",      "interview_engine.py + voice_interview.py")
    bold_kv(doc, "Technology",  "Google Gemini, AssemblyAI, gTTS, RapidFuzz, Python exec()")
    bold_kv(doc, "Input",       "Candidate resume data + job description context + candidate's text/voice answers")
    bold_kv(doc, "Output",      "Per-question scores, evaluation text, combined Stage 1+2 score, hire verdict")

    # RIGHT SIDE
    p = doc.add_paragraph()
    shade_para(p, 'FEF2F2')
    r = p.add_run("  RIGHT SIDE — SECURITY & VERIFICATION  ")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x7F, 0x1D, 0x1D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)

    sub_heading(doc, "Security Component 4D — Identity Verification / KYC")
    body(doc,
        "This component corresponds to identity_verification.py. It runs three independent "
        "cross-checks to confirm that the person submitting the resume is genuinely who they "
        "claim to be. "
        "NAME MATCH: Python difflib.SequenceMatcher compares the NER-extracted name from the "
        "resume against the name entered at registration. Threshold is 85% similarity ratio. "
        "EMAIL/PHONE MATCH: Exact case-insensitive email match plus last-10-digit phone comparison "
        "(to handle +91 prefix variations). "
        "FACE RECOGNITION: The profile photo embedded in the submitted PDF is extracted as base64 "
        "during ingestion. At verification time, this is compared against a live webcam photo "
        "captured in the browser using the face_recognition library (dlib HOG + CNN encodings). "
        "Euclidean face distance threshold = 0.6. If the distance is below 0.6, the faces match. "
        "The KYC label in the diagram stands for Know Your Candidate — all 4 checks must pass. "
        "OTP verification (Resend API email + timestamp expiry) acts as the final gate."
    )
    bold_kv(doc, "Module",      "identity_verification.py")
    bold_kv(doc, "Technology",  "face_recognition (dlib), difflib, Resend API, SQLite")
    bold_kv(doc, "Input",       "Resume photo (base64) + webcam snapshot + registered user details")
    bold_kv(doc, "Output",      "Verification result: name_match, email_match, phone_match, face_match (True/False + scores)")

    sub_heading(doc, "Security Component 4E — Fraud Detection")
    body(doc,
        "This component corresponds to fraud_detector.py. It runs four parallel heuristic checks "
        "on every resume, entirely without AI or external APIs — purely rule-based and explainable. "
        "CHECK 1 — Skill Inconsistency: Every skill listed in the Skills section is searched "
        "using word-boundary regex in the experience, projects, and education text. If a skill has "
        "zero body mentions, it is flagged as 'listed but not demonstrated.' "
        "CHECK 2 — Keyword Stuffing: Word frequency analysis detects any non-stopword appearing "
        "more than 8 times with a density exceeding 3% of total resume text. "
        "CHECK 3 — Hidden Text: Detects if raw word count is 2.5x larger than parsed count "
        "(white text ATS manipulation), finds 5+ consecutive identical word repetitions, or "
        "detects wall-of-text lines with 20+ words but zero punctuation. "
        "CHECK 4 — Temporal Consistency: Extracts all date ranges from experience, detects "
        "impossible timelines (end < start), future start dates, tenures >15 years, and "
        "overlapping employment periods >3 months. "
        "Risk score = sum of weighted flags × 15 (capped at 100). "
        "Risk levels: Low (<=20), Medium (<=55), High (>55). "
        "On Medium/High, Gemini optionally provides a natural language verdict."
    )
    bold_kv(doc, "Module",      "fraud_detector.py")
    bold_kv(doc, "Technology",  "Python regex, word frequency analysis, dateutil, optional Gemini LLM")
    bold_kv(doc, "Input",       "Cleaned resume text + raw resume word count + skills section")
    bold_kv(doc, "Output",      "fraud_risk (Low/Medium/High), risk_score (0-100), list of specific flags")

    sub_heading(doc, "Security Component 4F — Security & Access Control")
    body(doc,
        "This refers to the overall access control layer embedded across the application: "
        "Role-based routing ensures candidates cannot access recruiter dashboards and vice versa. "
        "Webcam monitoring (webcam_monitor.py / MediaPipe) runs as a continuous background thread "
        "during the entire interview — detecting face presence every 0.5 seconds. A 2-strike "
        "violation system terminates the interview if the candidate leaves the frame for more "
        "than 2 seconds or if multiple faces are detected. CLAHE (Contrast Limited Adaptive "
        "Histogram Equalization) is applied to the video frames to handle low-lighting "
        "conditions in remote settings. API keys (GEMINI_API_KEY, ASSEMBLYAI_API_KEY, "
        "RESEND_API_KEY) are loaded from a .env file via python-dotenv and never exposed "
        "in the source code."
    )
    bold_kv(doc, "Module",      "webcam_monitor.py, app.py (role guards), python-dotenv")
    bold_kv(doc, "Technology",  "MediaPipe Face Detection, streamlit-webrtc (WebRTC), CLAHE, python-dotenv")
    bold_kv(doc, "Input",       "Live webcam video stream + session role context")
    bold_kv(doc, "Output",      "Violation count, face status ('FACE OK'/'NO FACE'/'MULTIPLE FACES'), interview termination signal")

    divider(doc)

    # ── LAYER 5 ───────────────────────────────────────────────────────────────
    layer_banner(doc, 5, "DATA STORAGE", '581C87')

    body(doc,
        "This is the bottommost layer — the persistent storage tier. The diagram shows THREE distinct "
        "storage boxes in Layer 5, representing three different types of data storage used by the system. "
        "All three are accessed via Python code, with no external database server required."
    )

    sub_heading(doc, "Storage 5A — Relational Database (User Data, Profiles) — SQLite3")
    body(doc,
        "SQLite3 stores all structured, relational data: "
        "USERS TABLE — user_id, name, email, password_hash, role (recruiter/candidate), otp_code, otp_expiry, created_at. "
        "CANDIDATES TABLE — candidate_id, user_id (FK), name, email, skills_json, ats_score, "
        "fraud_risk, interview_score, voice_score, combined_score, hire_status, job_id (FK), resume_path, profile_photo_b64. "
        "INTERVIEW_ANSWERS TABLE — answer_id, candidate_id (FK), question_text, answer_text, "
        "score (0-10), evaluation_text, question_type (coding/technical/behavioral), mode (normal/ai/voice). "
        "SQLite was chosen because it requires zero server setup, is bundled with Python's stdlib, "
        "and is perfectly adequate for single-institution scale deployment."
    )
    bold_kv(doc, "Technology", "SQLite3 (Python stdlib)")
    bold_kv(doc, "Accessed by", "auth_db.py, candidate_db.py, interview_engine.py")

    sub_heading(doc, "Storage 5B — Object Storage (Resumes, Media) — Local Filesystem")
    body(doc,
        "All uploaded resume files are saved to a structured local directory: "
        "data/resumes/raw/ stores the original uploaded files with timestamped names to prevent "
        "collisions. The filename is sanitized using regex (replaces non-alphanumeric except dots/hyphens). "
        "data/resumes/processed/ would store cleaned text outputs. "
        "Profile photos extracted from PDFs are stored as base64 strings in the SQLite candidates "
        "table (not as separate files), making them self-contained within the database record."
    )
    bold_kv(doc, "Technology", "Python os, pathlib, base64 encoding")
    bold_kv(doc, "Accessed by", "data_ingestion.py (save_uploaded_file)")

    sub_heading(doc, "Storage 5C — Vector Database (Semantic Embeddings) — FAISS")
    body(doc,
        "The FAISS vector store persists all candidate resume embeddings between application restarts. "
        "It consists of two files: faiss.index — a binary FAISS IndexFlatIP file containing all "
        "resume chunk vectors (384 dimensions each). Each vector represents a semantic chunk "
        "(e.g., one work experience entry) of a candidate's resume. "
        "metadata.pkl — a Python-pickled list of dictionaries, one per vector, containing: "
        "chunk_text (the actual text), section (e.g., 'experience'), "
        "source (resume filename), name (candidate name), chunk_id (unique ID). "
        "Also saved: metadata.json (human-readable version for debugging). "
        "When a new candidate is added, their old vectors are identified by name and removed, "
        "new chunks are appended, and the full index is rebuilt from scratch. This ensures "
        "deduplication across re-uploads."
    )
    bold_kv(doc, "Technology", "FAISS (Meta AI), Sentence Transformers (all-MiniLM-L6-v2), pickle, JSON")
    bold_kv(doc, "Accessed by", "embeddings.py, similarity.py")
    bold_kv(doc, "Index Type",  "IndexFlatIP (Inner Product = cosine similarity when vectors are L2-normalized)")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  PART B — STEP-BY-STEP FLOW
    # ══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    shade_para(p, '0F172A')
    r = p.add_run("  PART B: STEP-BY-STEP DATA FLOW THROUGH THE ARCHITECTURE  ")
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    body(doc,
        "This section traces the journey of a candidate through each layer of the architecture "
        "from first interaction to final hire decision, with Input → Processing → Output for each step."
    )

    # STEP 1
    step_banner(doc, 1, "JOB CREATION (Recruiter Action — Layers 1 & 3)")
    iop_table(doc,
        inp  = "Recruiter logs in via Web Dashboard. Fills in: job title, department, description text, required skills list.",
        proc = "Layer 1 (Web Dashboard) collects the form data. Layer 2 (Auth) validates the recruiter session. "
               "Layer 3 Job Service stores the job record in SQLite with a unique Job ID. "
               "Required skills are parsed and stored as a JSON list field for later matching.",
        out  = "Job record in SQLite (job_id, title, department, description, required_skills_json). "
               "Job is now visible in the candidate portal and ready to receive applications."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "The Job ID is a foreign key that links every candidate's application to a specific opening. "
        "The required_skills field is later used by similarity.py to compute skill match ratios "
        "between the job and each candidate's extracted skills. Without this step, the matching "
        "engine has no target to match against."
    )
    speaking_box(doc,
        "The recruiter starts by creating a job post on the Web Dashboard — our Streamlit frontend. "
        "They fill in the job title, description, and the required skills. This job is stored in "
        "our SQLite database with a unique ID. Every candidate who applies is linked to this job ID, "
        "which is what enables targeted semantic ranking later."
    )

    # STEP 2
    step_banner(doc, 2, "RESUME UPLOAD (Candidate Action — Layers 1, 2, 3)")
    iop_table(doc,
        inp  = "Candidate registers, verifies OTP, selects job opening, and uploads resume file (PDF/DOCX/TXT).",
        proc = "Layer 2 (Auth Gateway) validates the candidate session and OTP status. "
               "Layer 3 (Candidate Service) receives the uploaded file via Streamlit UploadedFile. "
               "data_ingestion.save_uploaded_file() sanitizes the filename using regex and saves "
               "the raw file to data/resumes/raw/ with a timestamp prefix to prevent collisions.",
        out  = "Raw resume file saved to local filesystem. Path stored in session state. "
               "Candidate record created in SQLite with pending analysis status."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "The file must be safely saved to disk before text extraction begins. The timestamp in "
        "the filename prevents overwrite conflicts if the same candidate re-uploads. The file path "
        "is stored in the SQLite candidates table so the recruiter can retrieve it later for manual review."
    )
    speaking_box(doc,
        "The candidate registers and verifies their email via OTP. Then they select the job they "
        "are applying for and upload their resume — PDF, DOCX, or TXT are all supported. Our "
        "data ingestion module sanitizes and saves the file securely before passing it to the "
        "processing pipeline."
    )

    # STEP 3
    step_banner(doc, 3, "RESUME PARSING & NLP EXTRACTION (Layer 4A — AI Modules)")
    iop_table(doc,
        inp  = "Saved raw resume file (PDF/DOCX/TXT).",
        proc = "data_ingestion.py: PyPDF2 (primary) or pdfminer.six (fallback) extracts raw text. "
               "Page 1 XObject images are scanned to extract the profile photo (largest image by area). "
               "clean_text() applies 7 normalization steps: re-encode UTF-8, remove control chars, "
               "normalize unicode punctuation, collapse spaces, collapse blank lines, strip, final strip. "
               "text_processing.build_structured_json(): "
               "(1) Advanced clean — remove boilerplate, URLs, punctuation-only lines. "
               "(2) Section segmentation — heading regex splits text into labeled sections. "
               "(3) Entity extraction — spaCy NER (name, location) + regex (email, phone). "
               "(4) Skill extraction — 200+ taxonomy match on 'skills' section only. "
               "(5) JSON assembly into structured dict.",
        out  = "Structured JSON dict: {name, email, phone, location, skills[], summary, education, "
               "experience, projects, certifications, profile_photo_base64}. Stored in SQLite."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "This structured JSON is the single source of truth for every downstream operation. "
        "The fraud detector reads the skills list. The ATS scorer reads the section keys. "
        "The embeddings module reads all the section text. The identity verifier reads the name "
        "and email. The interview engine reads the skills and projects. If parsing is poor, "
        "every downstream module suffers — so robust parsing is critical."
    )
    speaking_box(doc,
        "Once uploaded, the resume goes through our NLP parsing pipeline. We extract text using "
        "PyPDF2 with pdfminer as a fallback. Then spaCy's Named Entity Recognition identifies "
        "the candidate's name and location. Regex extracts the email and phone. A section "
        "segmentation algorithm splits the resume into labeled buckets — Skills, Experience, "
        "Education, Projects — and we extract 200+ skills using our custom taxonomy. "
        "All of this becomes a structured JSON object that every other module uses."
    )

    # STEP 4
    step_banner(doc, 4, "VECTOR EMBEDDING & FAISS STORAGE (Layer 4B + Layer 5C)")
    iop_table(doc,
        inp  = "Structured resume JSON from Step 3.",
        proc = "embeddings.prepare_text_chunks(): Each section is converted to a labeled text block "
               "(e.g., 'Work Experience:\\nBuilt ML pipeline...'). An overview chunk combines name + top skills + summary. "
               "chunk_text(): LangChain RecursiveCharacterTextSplitter splits long blocks (chunk_size=400, overlap=80). "
               "generate_embeddings(): all-MiniLM-L6-v2 encodes all chunks → L2-normalized float32 (n_chunks × 384). "
               "If a FAISS index exists, old chunks for this candidate are removed (name-based filter). "
               "New vectors appended → full index rebuilt. "
               "save_vector_store(): faiss.index and metadata.pkl saved to disk.",
        out  = "Candidate's resume vectors stored in persistent FAISS IndexFlatIP. "
               "metadata.pkl updated with chunk_text, section, source, name per vector."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "Semantic embedding transforms text meaning into mathematics. Two resumes that say "
        "'developed NLP pipeline' and 'built language model system' will produce similar "
        "vectors even without sharing keywords — this is what makes HireMind AI superior "
        "to traditional ATS keyword matching systems. The FAISS index enables searching "
        "across all candidates in milliseconds even as the database grows."
    )
    speaking_box(doc,
        "After parsing, we convert the resume into mathematical vectors using Sentence Transformers — "
        "specifically the all-MiniLM-L6-v2 model that produces 384-dimensional embeddings. "
        "These embeddings capture the semantic meaning of each resume section. We store them "
        "in a FAISS vector database — an ultra-fast similarity search library developed by "
        "Meta AI. This is what enables us to match resumes to job descriptions based on "
        "MEANING, not just keywords."
    )

    # STEP 5
    step_banner(doc, 5, "ATS SCORING + FRAUD DETECTION (Layer 4A & 4E — Parallel)")
    iop_table(doc,
        inp  = "Cleaned resume text + structured JSON from Step 3. Runs SIMULTANEOUSLY in parallel.",
        proc = "ATS SCORING (ats_scorer.py): "
               "check_section_presence() — awards 25 points each for Skills, Experience, Education, Projects present. "
               "validate_experience_structure() — checks for date presence (+50), chronological order (+30), content depth (+20). "
               "Final ATS = (section_score × 0.60) + (structure_score × 0.40). "
               "Grade assigned: >=80 Excellent, >=60 Good, >=40 Fair, <40 Poor. "
               "FRAUD DETECTION (fraud_detector.py): "
               "Skill Inconsistency check (weight 1) + Keyword Stuffing check (weight 2) + "
               "Hidden Text check (weight 3) + Temporal Consistency check (weight 3). "
               "risk_score = sum_of_weights × 15 (capped 100). "
               "Risk level: Low <=20, Medium <=55, High >55. "
               "On Medium/High: optional Gemini LLM validation called.",
        out  = "ATS score (0-100) + grade + actionable feedback list. "
               "Fraud risk level (Low/Medium/High) + risk_score + flag details. "
               "All stored in SQLite candidates table."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "ATS scoring gives recruiters an immediate, explainable signal about resume quality — "
        "without AI bias. Fraud detection protects the pipeline from the growing problem of "
        "AI-generated, keyword-stuffed resumes that fool traditional systems. Both run "
        "deterministically — no LLM required — making them fast, reliable, and auditable. "
        "A candidate with a High fraud risk can be deprioritized before any interview resources are spent."
    )
    speaking_box(doc,
        "While embedding is happening, two other checks run in parallel. First, our ATS scoring "
        "engine — it checks whether the resume has all four required sections, whether dates "
        "are present and in chronological order, and scores it out of 100. Second, our fraud "
        "detection module runs four independent checks — does the candidate actually "
        "demonstrate their listed skills? Is any keyword appearing suspiciously often? "
        "Are there hidden text anomalies? Are employment dates consistent? This gives us a "
        "Low, Medium, or High fraud risk score before we ever speak to the candidate."
    )

    # STEP 6
    step_banner(doc, 6, "SEMANTIC MATCHING & CANDIDATE RANKING (Layer 4B — Matching)")
    iop_table(doc,
        inp  = "Job Description text (typed by recruiter) + FAISS index of all candidates (from Step 4).",
        proc = "similarity.process_job_description(): clean JD text + extract required skills using same taxonomy. "
               "generate_jd_embedding(): encode JD with all-MiniLM-L6-v2 + L2-normalize → 384-dim vector. "
               "FAISS IndexFlatIP.search(): returns top k=20 most similar resume chunks with cosine scores. "
               "aggregate_by_candidate(): group chunks by candidate name, compute max_score and mean_score, "
               "combined_semantic = 0.6×max + 0.4×mean. Filter chunks below 0.40 threshold. "
               "filter_by_hard_skills(): case-insensitive set intersection → matched_skills, missing_skills, match_ratio. "
               "rank_candidates(): final_score = (semantic × 0.65) + (skill_ratio × 0.35). Sort descending.",
        out  = "Ranked list of candidates with: name, final_score, semantic_score, skill_score, "
               "matched_skills[], missing_skills[]. Displayed on recruiter dashboard sorted by final_score."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "This is the central value proposition of HireMind AI. A candidate who built a "
        "'natural language question answering system' scores highly for an NLP role even "
        "if they didn't write the exact phrase 'NLP.' Traditional ATS systems would miss "
        "this candidate. The dual weighting (semantic 65% + skill overlap 35%) ensures "
        "both conceptual relevance AND specific technical alignment."
    )
    speaking_box(doc,
        "When the recruiter enters the job description, we embed it into the same vector space "
        "as all the resumes. Then FAISS finds the most semantically similar resume chunks across "
        "all candidates — in milliseconds. We aggregate these chunk scores by candidate, apply "
        "a skill match filter, and produce a final ranking. The formula is 65% semantic "
        "similarity plus 35% skill keyword overlap. The recruiter sees an instant ranked "
        "leaderboard of candidates, sorted by how well they match the role."
    )

    # STEP 7
    step_banner(doc, 7, "IDENTITY VERIFICATION — KYC GATE (Layer 4D — Security)")
    iop_table(doc,
        inp  = "Candidate's resume (name, email, profile_photo_base64) + registration details + OTP + live webcam photo.",
        proc = "Name check: difflib.SequenceMatcher >= 85% between NER name and registered name. "
               "Email check: exact case-insensitive match. "
               "Phone check: last-10-digit suffix comparison. "
               "OTP check: 6-digit code match + time.time() expiry validation (sent via Resend API). "
               "Face check: decode base64 resume photo → face_recognition.face_encodings() "
               "→ decode webcam snapshot → face_encodings() → compare_faces(tolerance=0.6) "
               "→ face_distance() → similarity_score = 1 - distance. "
               "All 4 checks evaluated; result returned as dict with per-check status and scores.",
        out  = "Verification result: {name_match: True/False, email_match: True/False, "
               "phone_match: True/False, face_match: True/False, face_similarity: 0.0-1.0}. "
               "Only fully verified candidates proceed to the interview stage."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "Without identity verification, there is nothing stopping a candidate from submitting "
        "a borrowed resume or having someone else take the interview for them. The face recognition "
        "step is particularly powerful — it compares the photo embedded in the submitted PDF "
        "(which the candidate cannot easily change retrospectively) against a live camera capture. "
        "This directly prevents impersonation fraud."
    )
    speaking_box(doc,
        "Before the interview begins, every candidate must pass our four-layer identity verification. "
        "We fuzzy-match their name, exactly match their email, compare phone number suffixes, "
        "verify a time-limited OTP sent to their email, AND compare their resume photo against "
        "a live webcam capture using facial recognition with a dlib-based 128-point face encoding. "
        "All four checks must pass. This is our KYC — Know Your Candidate layer."
    )

    # STEP 8
    step_banner(doc, 8, "AI INTERVIEW SYSTEM — TEXT + VOICE (Layer 4C)")
    iop_table(doc,
        inp  = "Verified candidate + structured resume JSON (skills, projects) + job description context.",
        proc = "TEXT INTERVIEW (Stage 1 — interview_engine.py): "
               "Gemini generates 7 personalized questions (3 coding + 2 technical + 2 behavioral). "
               "Coding: safe exec() + test cases + RapidFuzz similarity → composite score. "
               "Text: RapidFuzz token_sort_ratio (70%) + keyword boost (30%). "
               "Each question scored 0-10. Stage 1 score = average across 7 questions as percentage. "
               "MediaPipe webcam monitor active: 2-strike violation system runs in background thread. "
               "VOICE INTERVIEW (Stage 2 — voice_interview.py): "
               "Gemini generates 3 spoken questions (project/skill/scenario). "
               "gTTS converts to MP3 → st.audio() plays in browser. "
               "Candidate speaks → audio uploaded to AssemblyAI → Universal-3-pro transcribes → "
               "Gemini evaluates transcript (correctness, depth, relevance, communication) → score 0-10. "
               "Combined: final_pct = (stage1_pct × 0.60) + (stage2_pct × 0.40). "
               "Verdict: >=75% Strong Hire, >=55% Hire, <55% No Hire.",
        out  = "Per-question scores + evaluation text (all 10 questions). "
               "Stage 1 score, Stage 2 score, Combined score, Hire verdict. "
               "Full Gemini synthesis report: strengths, weaknesses, skill gaps, recommendation. "
               "Stored in SQLite interview_answers table + candidates table."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "Traditional interviews rely on the availability and subjectivity of a human interviewer. "
        "HireMind AI's interview system runs 24/7, evaluates every candidate on exactly the same "
        "criteria, generates questions tailored to THEIR resume rather than a generic template, "
        "and provides a written evaluation that the recruiter can review. The voice component "
        "additionally assesses communication skills — critical for most roles — without the "
        "scheduling overhead of a phone screen."
    )
    speaking_box(doc,
        "The interview system has two stages. In Stage 1, Gemini generates 7 custom questions "
        "based on the candidate's own resume — coding questions tested with real code execution, "
        "technical and behavioral questions scored with fuzzy matching. Throughout this, our "
        "webcam monitor is watching using MediaPipe — two violations terminate the interview. "
        "In Stage 2, a voice interview: questions are read aloud via gTTS, the candidate "
        "speaks their answers, AssemblyAI transcribes the speech, and Gemini evaluates. "
        "The final score combines Stage 1 at 60% and Stage 2 at 40%, producing a verdict: "
        "Strong Hire, Hire, or No Hire."
    )

    # STEP 9 (FINAL)
    step_banner(doc, 9, "FINAL EVALUATION & ADMIN DASHBOARD (Layer 1 — Recruiter View)")
    iop_table(doc,
        inp  = "All candidate scores from SQLite: ATS score, fraud risk, interview score, voice score, combined score.",
        proc = "Admin Dashboard (Streamlit recruiter view) pulls all candidate records for a specific Job ID "
               "from SQLite. Displays ranked table sorted by combined_score (ATS + interview weighted). "
               "Recruiter can: click 'View Report' to see full per-candidate breakdown including "
               "per-question scores, Gemini evaluation, fraud flags, and hire verdict. "
               "Click 'Select Candidate' → Resend API sends automated hiring email → "
               "SQLite updates hire_status to 'hired'. "
               "Recommendation Engine also runs: identify_skill_gaps() + suggest_alternative_roles() "
               "generates a personalized learning path for each candidate regardless of outcome.",
        out  = "Recruiter makes data-driven hiring decision. Selected candidate receives automated "
               "hiring email. Rejected candidates optionally receive a skill gap report with "
               "learning recommendations (courses, projects, certifications)."
    )
    sub_heading(doc, "Why this step matters:")
    body(doc,
        "Every upstream step — parsing, embedding, fraud detection, ATS scoring, identity verification, "
        "interview — feeds into this final view. The recruiter sees a single ranked table, backed "
        "by a comprehensive per-candidate report, and can make a hiring decision by clicking one "
        "button. The entire pipeline — from resume upload to hiring decision — can complete "
        "without a single human having to read a resume or conduct a manual interview."
    )
    speaking_box(doc,
        "Finally, everything feeds into the Admin Dashboard. The recruiter sees all candidates "
        "ranked by their combined score. They can expand any candidate to see the complete "
        "report — ATS score, fraud risk assessment, per-question interview feedback, and the "
        "Gemini hire verdict. One click sends the hiring email automatically through Resend API "
        "and marks the candidate as hired in our database. The entire hiring cycle — from "
        "resume upload to decision — is fully automated."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  PART C — QUICK PRESENTATION VERSION
    # ══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    shade_para(p, '0F172A')
    r = p.add_run("  PART C: QUICK PRESENTATION VERSION (60-90 SECOND SCRIPT)  ")
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p2 = doc.add_paragraph()
    shade_para(p2, 'DBEAFE')
    r2 = p2.add_run(
        "  USE THIS when a judge asks: 'Can you walk me through your system architecture diagram?'  "
        "Point to Figure 1 on the poster as you speak each layer.  "
    )
    r2.font.size = Pt(11); r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(14)

    script_paras = [
        ("LAYER 1 — User Interface",
         "At the very top of our architecture, we have the User Interface layer — built with Streamlit. "
         "On the left side, recruiters access the Web Dashboard to post jobs and view candidate rankings. "
         "On the right side, candidates access the Resume Upload Portal to submit their applications."),
        ("LAYER 2 — API Gateway & Authentication",
         "Below the UI, every single request passes through our API Gateway and Authentication layer. "
         "This handles OTP-based email verification using the Resend API, role-based access control "
         "to separate recruiter and candidate permissions, and session management via Streamlit's "
         "built-in session state — all backed by our SQLite database."),
        ("LAYER 3 — Backend Processing Services",
         "Layer 3 is our orchestration layer with four services: the Job Service stores job postings, "
         "the Candidate Service triggers the full analysis pipeline on resume upload, "
         "the Job Management Service coordinates the semantic matching phase, "
         "and the Workflow Orchestrator — our main app.py — coordinates everything together."),
        ("LAYER 4 (Left) — AI Modules & Processing",
         "The left half of Layer 4 contains our three core AI processing components. "
         "First, Resume Analysis and Parsing — using spaCy NER for entity extraction and our section "
         "segmentation algorithm. Second, Candidate Evaluation and Matching — using FAISS vector search "
         "and our ATS scoring engine. Third, the AI Interview System — personalized text and voice "
         "interviews powered by Google Gemini and AssemblyAI."),
        ("LAYER 4 (Right) — Security & Verification",
         "The right half of Layer 4 is our security perimeter. Identity Verification uses dlib face "
         "recognition to match the resume photo against a live webcam capture. Our Fraud Detection "
         "module runs four independent rule-based checks on every resume. And our Access Control "
         "includes continuous webcam proctoring with MediaPipe during the interview."),
        ("LAYER 5 — Data Storage",
         "At the base, our Data Storage layer has three components: SQLite stores all structured "
         "relational data — users, candidates, and interview answers. Local object storage holds "
         "the uploaded resume files. And our FAISS vector database holds the semantic embeddings "
         "of every resume, enabling millisecond similarity search across all candidates."),
        ("Summary",
         "What makes this architecture powerful is how all six layers work together seamlessly. "
         "A candidate uploads a resume at Layer 1. It flows down through authentication, "
         "through orchestration, through all the AI modules in parallel, and the results "
         "are persisted at Layer 5. When the recruiter needs results, they flow back up to "
         "Layer 1 as a ranked, scored, verified leaderboard. The entire pipeline runs "
         "autonomously. Thank you."),
    ]

    for label, text in script_paras:
        sub_heading(doc, f"[Point to {label} on diagram]", color=(126, 34, 206))
        p = doc.add_paragraph()
        shade_para(p, 'F5F3FF')
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(12)

    # ── FINAL COVER ───────────────────────────────────────────────────────────
    doc.add_page_break()
    ep = doc.add_paragraph()
    shade_para(ep, '0F172A')
    er = ep.add_run("  HireMind AI  |  System Architecture Explanation  |  CALIBRE 2k26  ")
    er.font.size = Pt(13); er.font.bold = True
    er.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(
        "Kartik Gupta  |  Bhavesh Chaudhary  |  Durga Muhto  |  Sudiksha Singh  |  Komal Yadav\n"
        "Thakur Shyamanarayan Engineering College, Mumbai\n"
        "18th April, 2026"
    )
    tr.font.size = Pt(11); tr.font.italic = True
    tr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # SAVE
    out_path = r"c:\Users\LENOVO\OneDrive\Desktop\HireMind_AI_System_Architecture_Explanation.docx"
    doc.save(out_path)
    print(f"[SUCCESS] Saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
